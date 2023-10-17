from docx import Document
from docx.shared import Pt, Mm
import os

def save_document(dataset, user_id):
    doc = Document('new_resume.docx')
    all_tables = doc.tables
    r = 0
    e = 0
    o = 0
    k = 0

    if dataset['russian'] == 'Basic':
        r = 1
    elif dataset['russian'] == 'Elementary':
        r = 2
    elif dataset['russian'] == 'Lower Intermediate':
        r = 3
    elif dataset['russian'] == 'Intermediate':
        r = 4
    elif dataset['russian'] == 'Upper Intermediate':
        r = 5
    elif dataset['russian'] == 'Advanced':
        r = 6
    else:
        r = 7


    if dataset['english'] == 'Basic':
        e = 1
    elif dataset['english'] == 'Elementary':
        e = 2
    elif dataset['english'] == 'Lower Intermediate':
        e = 3
    elif dataset['english'] == 'Intermediate':
        e = 4
    elif dataset['english'] == 'Upper Intermediate':
        e = 5
    elif dataset['english'] == 'Advanced':
        e = 6
    else:
        e = 7

    if 'other_lang_level' in dataset:
        if dataset['other_lang_level'] == 'Basic':
            o = 1
        elif dataset['other_lang_level'] == 'Elementary':
            o = 2
        elif dataset['other_lang_level'] == 'Lower Intermediate':
            o = 3
        elif dataset['other_lang_level'] == 'Intermediate':
            o = 4
        elif dataset['other_lang_level'] == 'Upper Intermediate':
            o = 5
        elif dataset['other_lang_level'] == 'Advanced':
            o = 6
        else:
            o = 7
    
    if dataset['know_about_as'] == 'Google':
        k = 1
    elif dataset['know_about_as'] == 'Facebook':
        k = 2
    elif dataset['know_about_as'] == 'Instagram':
        k = 3
    elif dataset['know_about_as'] == 'Vk':
        k = 4
    elif dataset['know_about_as'] == 'Friends':
        k = 5
    elif dataset['know_about_as'] == 'Other':
        k = 6
    else:
        k = 7

    temp = 1
    temp2 = 1
    temp3 = 1
    temp4 = 1

    for i, table in enumerate(all_tables):
        for j, row in enumerate(table.rows):
            for cell in row.cells:
                for word in cell.text.split():
                    if word == 'russian_test' and temp == r:
                        temp+=1
                        cell.text = cell.text.replace(word, 'Check')
                    elif word == 'russian_test':
                        temp+=1
                        cell.text = cell.text.replace(word, ' ')

                    if word == 'english_test' and temp2 == e:
                        temp2+=1
                        cell.text = cell.text.replace(word, 'Check')
                    elif word == 'english_test':
                        temp2+=1
                        cell.text = cell.text.replace(word, ' ')

                    if 'other_lang_level' in dataset:
                        if word == 'other_lang_level_test' and temp3 == o:
                            temp3+=1
                            cell.text = cell.text.replace(word, 'Check')
                        elif word == 'other_lang_level_test':
                            temp3+=1
                            cell.text = cell.text.replace(word, ' ')

                    if word == 'know_about_as_test' and temp4 == k:
                        temp4+=1
                        cell.text = cell.text.replace(word, 'Check')
                    elif word == 'know_about_as_test':
                        temp4+=1
                        cell.text = cell.text.replace(word, ' ')


    for i, table in enumerate(all_tables):
        for j, row in enumerate(table.rows):
            for cell in row.cells:
                for word in cell.text.split():
                    if cell.text == 'profile_photo':
                        if 'profile_photo' in word:
                            cell.text = cell.text.replace(word, '')
                        p = cell.paragraphs[0]
                        run = p.add_run()
                        run.add_picture(f'./{dataset["profile_Photo"]}', width=Mm(60))

    for i, table in enumerate(all_tables):
        cell = table.cell(0,0)
        for j, row in enumerate(table.rows):
            for cell in row.cells:

                for word in cell.text.split():
                    # print(i,'----', j, '--' ,word)                        
                    if '_test' in word:
                        if word.split('_test')[0] in dataset and dataset[word.split('_test')[0]]!=None :
                            cell.text = cell.text.replace(word, dataset[word.split('_test')[0]])
                            run = cell.paragraphs[0].runs[0]
                            font = run.font
                            font.name = 'Arial'
                            font.size = Pt(14)
    
    for i, table in enumerate(all_tables):
        for j, row in enumerate(table.rows):
            for cell in row.cells:
                for word in cell.text.split():
                    if '_test' in word:
                            cell.text = cell.text.replace(word, '')

    doc.add_paragraph()

    # Создаем параграф и добавляем текст
    imgP = doc.add_paragraph()
    imgP_run = imgP.add_run('Pictures:')
    imgP_run.bold = True
    imgP_run.font.size = Pt(14)
    imgP_run.font.name = 'Arial'

    # os.remove(f'./{dataset["profile_Photo"]}') 
    
    # if 'COVID_Photo' in dataset:
    #     doc.add_picture(f'./{dataset["COVID_Photo"]}', height = Mm(100))
    #     os.remove(f'./{dataset["COVID_Photo"]}') 

    if 'Course_Photo' in dataset:
            doc.add_picture(f'./{dataset["Course_Photo"]}', height = Mm(100))
            os.remove(f'./{dataset["Course_Photo"]}') 

    if 'tattoo_Photo' in dataset:
        doc.add_picture(f'./{dataset["tattoo_Photo"]}', height = Mm(100))
        os.remove(f'./{dataset["tattoo_Photo"]}') 

    if 'more_Photo' in dataset:
        doc.add_picture(f'./{dataset["more_Photo"]}', height = Mm(100))
        os.remove(f'./{dataset["more_Photo"]}') 
    
    if 'more_Photo2' in dataset:
        doc.add_picture(f'./{dataset["more_Photo2"]}', height = Mm(100))
        os.remove(f'./{dataset["more_Photo2"]}') 

    if 'more_Photo3' in dataset:
        doc.add_picture(f'./{dataset["more_Photo3"]}', height = Mm(100))
        os.remove(f'./{dataset["more_Photo3"]}') 

    if 'more_Photo4' in dataset:
        doc.add_picture(f'./{dataset["more_Photo4"]}', height = Mm(100))
        os.remove(f'./{dataset["more_Photo4"]}') 


    doc.save(f'{dataset["full_name"]}_{user_id}.docx')
    # doc.save('Check1.docx')
    return(f'{dataset["full_name"]}_{user_id}.docx')


dataset = {'language': 'ru', 'desired_positions': 'info 1', 'full_name': 'info 2', 
           'profile_Photo': 'userquest.jpg', 
           'date_of_birth': 'info 3', 'place_of_birth': 'info 4', 'married': 'N', 'have_any_children': 'N', 
           'height': 'info 5', 'weight': 'info 6', 'city_of_residence': 'info 7', 'current_location': 'info 8', 
           'nationality': 'info 9', 'nationality_country': 'info 10', 'haveP_': 'N', 'covid_access': 'Yes', 
           'COVID_Photo_NULL': 'AgACAgIAAxkBAAIH82US2bLh34iouz1SA48yOjzH9UAWAAK-0DEbvpqYSD_7xjd8eQZRAQADAgADeQADMAQ.jpg', 
           'phone_Number': 'info 11', 'messenger': 'info 12', 'email': 'info 13', 'instagram': 'info 14', 
           'Facebook': 'info 15', 'linkedIn': 'info 16', 'vkontakte': 'info 17', 'name_cousen': 'info 18', 
           'phone_Number_cousen': 'info 19', 'relative_cousen': 'info 20', 'father_name': 'info 21', 
           'mother_name': 'info 22', 'gatar': 'no', 'UAE': 'yes', 'Bahrain': 'no', 'Oman': 'yes', 
           'education_stepen': 'info 23', 'university_name': 'info 24', 'special_degree': 'info 25', 
           'year_of_Education': 'info 26', 'postgraduate_access': 'no', 'course_access': 'yes', 
           'сourse_name': 'info 27', 'course_date': 'info 28', 'Course_place': 'info 29', 
           'doc_course_access': 'yes', 'Course_Photo_NULL': 'AgACAgIAAxkBAAIIHWUS2fWSOI0uaQPpq-iSFpW000HSAALe0DEbvpqYSAclBDun5lnAAQADAgADbQADMAQ.jpg', 
           'tattoo_access': 'no', 'work_exp': 'info 30', 'work_name': 'info 31', 'work_place': 'info 32', 
           'work_position': 'info 33', 'work_responsibilities': 'info 34', 'other_work_access': 'no', 
           'now_work_access': 'yes', 'now_work_exp': 'info 35', 'now_work_name': 'info 36', 'now_work_place': 'info 37',
             'now_work_position': 'info 38', 'now_work_responsibilities': 'info 39', 'hoste_program': 'info 40', 
             'finance_program': 'info 41', 'travel_program': 'info 42', 'graph_program': 'info 43', 'car_access': 'no',
               'russian': 'Fluent', 'english': 'Upper Intermediate', 'other_lang_access': 'Yes', 
               'added_language': 'info 44', 'other_lang_level': 'Fluent', 'know_about_as': 'Friends',
               'msg_date_test':"26 September"}

save_document(dataset, 1231)
