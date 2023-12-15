from docx import Document
from docx.shared import Pt, Mm
import os

def save_document2(dataset, user_id):
    doc = Document('sample_resume2.docx')
    all_tables = doc.tables

    for i, table in enumerate(all_tables):
        for j, row in enumerate(table.rows):
            for cell in row.cells:
                for word in cell.text.split():
                    if cell.text == 'profile_photo':
                        if 'profile_photo' in word:
                            cell.text = cell.text.replace(word, '')
                        p = cell.paragraphs[0]
                        run = p.add_run()
                        run.add_picture(f'./{dataset["profile_Photo"]}', width=Mm(60))
                        # run.add_picture(f'./test_picture.jpg', width=Mm(60))

    for i, table in enumerate(all_tables):
        cell = table.cell(0,0)
        for j, row in enumerate(table.rows):
            for cell in row.cells:

                for word in cell.text.split():
                    # print(i,'----', j, '--' ,word)                        
                    if '_test' in word:
                        if word.split('_test')[0] in dataset and dataset[word.split('_test')[0]]!=None :
                            cell.text = cell.text.replace(word, dataset[word.split('_test')[0]])
                            run = cell.paragraphs[0].runs[0]
                            font = run.font
                            font.name = 'Century Gothic'
                            font.size = Pt(9)
    
    for i, table in enumerate(all_tables):
        for j, row in enumerate(table.rows):
            for cell in row.cells:
                for word in cell.text.split():
                    if '_test' in word:
                            cell.text = cell.text.replace(word, '')

    doc.add_paragraph()

    # Создаем параграф и добавляем текст
    imgP = doc.add_paragraph()
    imgP_run = imgP.add_run('Pictures:')
    imgP_run.bold = True
    imgP_run.font.size = Pt(9)
    imgP_run.font.name = 'Century Gothic'

    os.remove(f'./{dataset["profile_Photo"]}') 
    
    # if 'COVID_Photo' in dataset:
    #     doc.add_picture(f'./{dataset["COVID_Photo"]}', height = Mm(100))
    #     os.remove(f'./{dataset["COVID_Photo"]}') 

    if 'Course_Photo' in dataset:
            doc.add_picture(f'./{dataset["Course_Photo"]}', height = Mm(100))
            os.remove(f'./{dataset["Course_Photo"]}') 

    if 'tattoo_Photo' in dataset:
        doc.add_picture(f'./{dataset["tattoo_Photo"]}', height = Mm(100))
        os.remove(f'./{dataset["tattoo_Photo"]}') 

    if 'more_Photo' in dataset:
        doc.add_picture(f'./{dataset["more_Photo"]}', height = Mm(100))
        os.remove(f'./{dataset["more_Photo"]}') 
    
    if 'more_Photo2' in dataset:
        doc.add_picture(f'./{dataset["more_Photo2"]}', height = Mm(100))
        os.remove(f'./{dataset["more_Photo2"]}') 

    if 'more_Photo3' in dataset:
        doc.add_picture(f'./{dataset["more_Photo3"]}', height = Mm(100))
        os.remove(f'./{dataset["more_Photo3"]}') 

    if 'more_Photo4' in dataset:
        doc.add_picture(f'./{dataset["more_Photo4"]}', height = Mm(100))
        os.remove(f'./{dataset["more_Photo4"]}') 


    doc.save(f'{dataset["full_name"]}_{user_id}_2.docx')
    # doc.save('Check1.docx')
    return(f'{dataset["full_name"]}_{user_id}_2.docx')


# dataset = {'language': 'ru', 'desired_positions': 'Певец. Музыкант.', 'full_name': 'Абылкаир Сагыныш Айтжанович', 
#            'profile_Photo': 'test_picture.jpg', 'y_birth': '1996год.', 'm_birth': 'Сентябрь', 'd_birth': '25 Сентября', 'c_birth': 'Казахстан', 
#            'c_or_v_birth': 'Павлодар.', 'married': 'Нет', 'have_any_children': 'Нет', 'height': '180см.', 'weight': '75кг', 'city_of_residence': 'Алматы.', 
#            'current_location': 'Алматы.', 'nationality': 'Казах.', 'nationality_country': 'Казахстан', 'haveP_': 'Да', 'covid_access': 'Yes', 
#            'phone_Number': '87711634232.', 'messenger': 'Да.', 'email': 'Saga.aa.96@mail.ru', 'instagram': 'sagynysh__abylkair', 'name_cousen': '87051739801\nЛяззат(мама)', 
#            'phone_Number_cousen': '87055537159', 'relative_cousen': 'Сестра.', 'father_name': 'Айтжан.', 'mother_name': 'Ляззат.', 'gatar': 'yes', 'UAE': 'yes', 'Bahrain': 'yes', 
#            'Oman': 'yes', 'education_stepen': 'Высшее. Степень бакалавра.', 'university_name': 'Колледж Казнам. Город Астана. Специальность Вокальное искусство. Бакалавр ПГПИ. Город Павлодар. Специальность Музыкальное образование.', 
#            'special_degree': 'Вокальное искусство.', 'year_of_Education': '2020-2022.', 'postgraduate_access': 'no', 'course_access': 'yes', 
#            'сourse_name': 'Курс Эстрадного вокала КазНАИ им. Жургенова и Академического вокала при Кнк им. Курмангазы в городе Алматы.', 'course_date': '2022-2023', 'Course_place': 'В городе Алматы.', 
#            'doc_course_access': 'no', 'tattoo_access': 'no', 'work_exp': '2022-2023.', 'work_name': 'Оперный театр им. Абая в городе Алматы.', 'work_place': 'Алматы.', 'work_position': 'Солист хора.', 
#            'work_responsibilities': 'Пение в коллективе.', 'other_work_access': 'yes', 'other_work_exp': '2019-2023.', 'other_work_name': 'Филармония города Алматы.', 'other_work_place': 'Алматы.', 
#            'other_work_position': 'Солист.', 'other_work_responsibilities': 'Выступление на концертах.', 'now_work_access': 'no', 'hoste_program': 'Нет', 'finance_program': 'Нет', 'travel_program': 'Нет', 
#            'graph_program': 'Нет', 'other_programs': 'Нет', 'car_access': 'yes', 'car_category': 'Б', 'russian': 'Advanced', 'english': 'Intermediate', 'other_lang_access': 'Yes', 'added_language': 'Казахский.',
#              'other_lang_level': 'Fluent', 'know_about_as': 'Friends', 
#             #  'more_Photo': 'AgACAgIAAxkBAAImgmV7NoOrD6jv43Dr6Y0kjUaYxhkNAAKA0TEbf0TYS7C-npLPtXCSAQADAgADeQADMwQ.jpg', 
#             #  'more_Photo2': 'AgACAgIAAxkBAAImhGV7Nqjr0NPm5sFO05JeZMdK_oyoAAKB0TEbf0TYS8-WT7MTMki1AQADAgADeQADMwQ.jpg', 
#             #  'more_Photo3': 'AgACAgIAAxkBAAImhmV7Nrh09JIODSTix-KRw61HF5zLAAKC0TEbf0TYS3tLpGPLLbv7AQADAgADeQADMwQ.jpg', 
#             #  'more_Photo4': 'AgACAgIAAxkBAAImiGV7NvMdjP3BB5n1rzQYgXOrB38SAAKD0TEbf0TYS5wGUConDW0QAQADAgADeQADMwQ.jpg', 
#              'srok_job': 'Год.', 'msg_date': '14 December 17:10', 'resume_creating_date': '14 December 17:10'}

# save_document(dataset, 1231)
