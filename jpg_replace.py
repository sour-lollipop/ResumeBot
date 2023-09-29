from docx.oxml import parse_xml, register_element_cls
from docx.oxml.ns import nsdecls
from docx.oxml.shape import CT_Picture
from docx.oxml.xmlchemy import BaseOxmlElement, OneAndOnlyOne

# смотрите: docx.oxml.shape.CT_Inline
class CT_Anchor(BaseOxmlElement):
    """
    Элемент `<w:anchor>`, контейнер для плавающего изображения.
    """
    extent = OneAndOnlyOne('wp:extent')
    docPr = OneAndOnlyOne('wp:docPr')
    graphic = OneAndOnlyOne('a:graphic')

    @classmethod
    def _anchor_xml(cls, pos_x, pos_y):
        """
        Стиль переноса текста: `<wp:anchor behindDoc="0">`;
        Положение изображения: `<wp:positionH relativeFrom="page">`;
        Обтекание текста: `<wp:wrapSquare wrapText="largest"/>`.
        """
        return (
            '<wp:anchor behindDoc="0" distT="0" distB="0" distL="0" distR="0"'
            ' simplePos="0" layoutInCell="1" allowOverlap="1" relativeHeight="2"'
            f' {nsdecls("wp", "a", "pic", "r")}>'
            '  <wp:simplePos x="0" y="0"/>'
            '  <wp:positionH relativeFrom="page">'
            f'    <wp:posOffset>{int(pos_x)}</wp:posOffset>'
            '  </wp:positionH>'
            '  <wp:positionV relativeFrom="page">'
            f'    <wp:posOffset>{int(pos_y)}</wp:posOffset>'
            '  </wp:positionV>'
            '  <wp:extent />'
            '  <wp:wrapSquare wrapText="largest"/>'
            '  <wp:docPr />'
            '  <wp:cNvGraphicFramePr>'
            '    <a:graphicFrameLocks noChangeAspect="1"/>'
            '  </wp:cNvGraphicFramePr>'
            '  <a:graphic>'
            '    <a:graphicData>'
            '    </a:graphicData>'
            '  </a:graphic>'
            '</wp:anchor>'
        )

    @classmethod
    def new(cls, cx, cy, shape_id, pic, pos_x, pos_y):
        """
        Возвращает новый элемент `<wp:anchor>`, заполненный 
        переданными значениями в качестве параметров.
        """
        anchor = parse_xml(cls._anchor_xml(pos_x, pos_y))
        anchor.extent.cx = cx
        anchor.extent.cy = cy
        anchor.docPr.id = shape_id
        anchor.docPr.name = f'Picture {shape_id}'
        anchor.graphic.graphicData.uri = (
                'http://schemas.openxmlformats.org/drawingml/2006/picture')
        anchor.graphic.graphicData._insert_pic(pic)
        return anchor

    @classmethod
    def new_pic_anchor(cls, shape_id, rId, filename, cx, cy, pos_x, pos_y):
        """
        Возвращает новый элемент `wp:anchor`, содержащий элемент 
        `pic:pic` задается значениями аргументов.
        """
        pic_id = 0  # Word, похоже, не использует это, но и не опускает его
        pic = CT_Picture.new(pic_id, filename, rId, cx, cy)
        anchor = cls.new(cx, cy, shape_id, pic, pos_x, pos_y)
        anchor.graphic.graphicData._insert_pic(pic)
        return anchor


# смотрите: docx.parts.story.BaseStoryPart.new_pic_inline
def new_pic_anchor(part, image_descriptor, width, height, pos_x, pos_y):
    """
    Возвращает вновь созданный элемент `w:anchor`.
    Элемент содержит изображение, указанное в *image_descriptor*,
    и масштабируется на основе значений *width* и *height*.
    """
    rId, image = part.get_or_add_image(image_descriptor)
    cx, cy = image.scaled_dimensions(width, height)
    shape_id, filename = part.next_id, image.filename    
    return CT_Anchor.new_pic_anchor(shape_id, rId, filename, cx, cy, pos_x, pos_y)

# смотрите: docx.text.run.add_picture
def add_float_picture(p, image_path_or_stream, width=None, height=None, pos_x=0, pos_y=0):
    """
    Добавляет плавающее изображение в фиксированном 
    положении "pos_x" и "pos_y", отсчет - левый верхний угол.
    """
    run = p.add_run()
    anchor = new_pic_anchor(run.part, image_path_or_stream, width, height, pos_x, pos_y)
    run._r.add_drawing(anchor)

# смотрите: docx.oxml.shape.__init__.py
register_element_cls('wp:anchor', CT_Anchor)


if __name__ == '__main__':

    from docx import Document
    from docx.shared import Mm

    doc = Document()
    # добавим плавающее изображение
    p = doc.add_paragraph()
    add_float_picture(p, '/path/to/image.jpg', width=Mm(25), pos_x=Mm(30), pos_y=Mm(30))
    # добавим текст
    p.add_run('текст документа. ' * 50)
    doc.save('test.docx')